import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_latex(text):
    # Remove comments
    text = re.sub(r'(?<!\\)%.*', '', text)
    # Remove common commands but keep text
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\uppercase\{([^}]*)\}', lambda m: m.group(1).upper(), text)
    text = re.sub(r'\\fontsize\{[^}]*\}\{[^}]*\}\\selectfont', '', text)
    text = re.sub(r'\\bfseries', '', text)
    text = re.sub(r'\\itshape', '', text)
    text = re.sub(r'\\normalfont', '', text)
    # Remove citations and labels
    text = re.sub(r'\\cite\{[^}]*\}', '[Ref]', text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    text = re.sub(r'\\ref\{[^}]*\}', '[Ref]', text)
    # Basic math cleanup
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    # Clean escapes
    text = text.replace('\\&', '&').replace('\\%', '%').replace('\\_', '_').replace('\\{', '{').replace('\\}', '}')
    return text.strip()

def parse_table(table_content):
    rows = []
    # Remove commands inside table
    table_content = re.sub(r'\\toprule|\\midrule|\\bottomrule|\\hline', '', table_content)
    # Split by \\
    lines = table_content.split('\\\\')
    for line in lines:
        if line.strip():
            # Split by &
            cols = [clean_latex(c.strip()) for c in line.split('&')]
            rows.append(cols)
    return rows

def convert():
    tex_path = 'MindTrace_Survey_IEEE.tex'
    docx_path = 'MindTrace_Survey_IEEE_Generated.docx'
    
    if not os.path.exists(tex_path):
        print(f"Error: {tex_path} not found.")
        return

    with open(tex_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Title
    title_match = re.search(r'\\title\{([^}]*)\}', content, re.DOTALL)
    if title_match:
        title_text = clean_latex(title_match.group(1))
        p = doc.add_heading(title_text, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author
    author_match = re.search(r'\\author\{(.*?)\}', content, re.DOTALL)
    if author_match:
        author_text = re.sub(r'\\IEEEauthorblockN\{([^}]*)\}', r'\1', author_match.group(1), flags=re.DOTALL)
        author_text = re.sub(r'\\IEEEauthorblockA\{([^}]*)\}', r'\n\1', author_text, flags=re.DOTALL)
        author_text = clean_latex(author_text)
        p = doc.add_paragraph(author_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', content, re.DOTALL)
    if abstract_match:
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(clean_latex(abstract_match.group(1)))

    # Main Content
    # Split by sections and subsections
    parts = re.split(r'(\\section\{[^}]*\}|\\subsection\{[^}]*\}|\\subsubsection\{[^}]*\})', content)
    
    for i in range(len(parts)):
        part = parts[i]
        if part.startswith('\\section{'):
            title = part[9:-1]
            doc.add_heading(clean_latex(title), level=1)
        elif part.startswith('\\subsection{'):
            title = part[12:-1]
            doc.add_heading(clean_latex(title), level=2)
        elif part.startswith('\\subsubsection{'):
            title = part[15:-1]
            doc.add_heading(clean_latex(title), level=3)
        else:
            # Process text and tables
            # Find tables
            table_parts = re.split(r'(\\begin\{tabular[x]?\}.*?\\end\{tabular[x]?\})', part, flags=re.DOTALL)
            for tp in table_parts:
                if tp.startswith('\\begin{tabular'):
                    # Extract rows
                    inner = re.search(r'\\begin\{tabular[x]?\}(?:\{[^}]*\}){1,2}(.*?)\\end\{tabular[x]?\}', tp, re.DOTALL)
                    if inner:
                        rows = parse_table(inner.group(1))
                        if rows:
                            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                            table.style = 'Table Grid'
                            for r_idx, row_data in enumerate(rows):
                                for c_idx, val in enumerate(row_data):
                                    if c_idx < len(table.rows[r_idx].cells):
                                        table.rows[r_idx].cells[c_idx].text = val
                else:
                    # Normal text - split into paragraphs
                    # Handle lists first
                    list_parts = re.split(r'(\\begin\{enumerate\}.*?\\end\{enumerate\}|\\begin\{itemize\}.*?\\end\{itemize\})', tp, flags=re.DOTALL)
                    for lp in list_parts:
                        if lp.startswith('\\begin{enumerate}') or lp.startswith('\\begin{itemize}'):
                            items = re.findall(r'\\item\s+(.*?)(?=\\item|\\end)', lp, re.DOTALL)
                            for item in items:
                                doc.add_paragraph(clean_latex(item), style='List Bullet' if 'itemize' in lp else 'List Number')
                        else:
                            paragraphs = lp.split('\n\n')
                            for p_text in paragraphs:
                                p_text = clean_latex(p_text)
                                if p_text:
                                    doc.add_paragraph(p_text)

    # References
    ref_match = re.search(r'\\begin\{enumerate\}\[label=\{.*?\}\](.*?)\\end\{enumerate\}', content, re.DOTALL)
    if ref_match:
        doc.add_heading('References', level=1)
        items = re.findall(r'\\item\s+(.*?)(?=\\item|\\end)', ref_match.group(1), re.DOTALL)
        for item in items:
            doc.add_paragraph(clean_latex(item), style='List Number')

    doc.save(docx_path)
    print(f"Success: {docx_path} created.")

if __name__ == "__main__":
    convert()
